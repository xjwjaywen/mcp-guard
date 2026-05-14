from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "docs" / "MCP-Guard实验室项目情况.docx"


def set_font(run, name="Calibri", east_asia="Microsoft YaHei", size=None, bold=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph(paragraph, after=6, before=0, line=1.10):
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.line_spacing = line


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph()
    set_paragraph(paragraph, before=16 if level == 1 else 12, after=8 if level == 1 else 6)
    run = paragraph.add_run(text)
    size = 16 if level == 1 else 13
    set_font(run, size=size, bold=True, color="2E74B5")
    return paragraph


def add_body(doc, text, bold_label=False):
    paragraph = doc.add_paragraph()
    set_paragraph(paragraph)
    if bold_label and "：" in text:
        label, body = text.split("：", 1)
        label_run = paragraph.add_run(label + "：")
        set_font(label_run, size=11, bold=True)
        body_run = paragraph.add_run(body)
        set_font(body_run, size=11)
    else:
        run = paragraph.add_run(text)
        set_font(run, size=11)
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    set_paragraph(paragraph, after=4, line=1.167)
    run = paragraph.add_run(text)
    set_font(run, size=11)
    return paragraph


def add_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    widths = [1800, 7560]
    for index, text in enumerate(("字段", "内容")):
        cell = table.rows[0].cells[index]
        set_cell_width(cell, widths[index])
        set_cell_shading(cell, "F2F4F7")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(text)
        set_font(run, size=10.5, bold=True)

    for label, value in rows:
        cells = table.add_row().cells
        for index, text in enumerate((label, value)):
            set_cell_width(cells[index], widths[index])
            cells[index].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraph = cells[index].paragraphs[0]
            set_paragraph(paragraph, after=2, line=1.10)
            run = paragraph.add_run(text)
            set_font(run, size=10.5, bold=index == 0)
    return table


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph(title, after=10)
    run = title.add_run("MCP-Guard 项目情况总结")
    set_font(run, size=24, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph(subtitle, after=18)
    run = subtitle.add_run("暨南大学网络空间安全学院")
    set_font(run, size=13, bold=True, color="555555")

    add_heading(doc, "MCP 工具调用安全审计器", 1)
    add_table(
        doc,
        [
            ("项目名称", "MCP 工具调用安全审计器"),
            ("英文名称", "MCP-Guard"),
            ("开源地址", "https://github.com/xjwjaywen/mcp-guard"),
            ("项目时间", "2026年5月14日—至今"),
            ("负责人", "待填写（网络空间安全专业，学生信息按实验室要求补充）"),
            ("指导教师", "耿光刚；刘东杰"),
            ("开发语言", "Python、HTML、CSS；支持 JSON、YAML、TOML 配置格式"),
            ("代码规模", "约 2635 行代码（不含虚拟环境、缓存和自动生成报告文件）"),
        ],
    )

    add_heading(doc, "项目内容", 1)
    add_body(
        doc,
        "MCP-Guard 是一个面向 AI Agent / MCP 工具调用生态的本地安全审计项目，"
        "用于扫描 Claude Desktop、Cursor、VS Code、Windsurf 等环境中的 MCP 配置文件，"
        "识别高风险启动命令、明文凭据、敏感目录暴露、过宽文件权限、网络暴露、Docker Socket 暴露、"
        "运行时包管理器启动和提示注入文本等风险，并生成结构化风险报告与修复建议。"
    )
    add_body(
        doc,
        "项目采用 Flask Web 仪表盘与 Python CLI 双入口设计，既可以在浏览器中上传、粘贴或自动发现配置文件，"
        "也可以通过 python -m mcp_guard scan 命令批量扫描本地配置。扫描结果会保存为 JSON 和 Markdown 文件，"
        "便于实验统计、项目汇报和后续规则扩展。"
    )

    add_heading(doc, "核心功能", 1)
    bullets = [
        "Web 仪表盘：支持上传、粘贴、示例扫描和本机常见 MCP 配置路径自动发现。",
        "CLI 工具：支持 python -m mcp_guard scan <配置文件> 和 discover --scan。",
        "多格式解析：支持 JSON、YAML、TOML 三类配置格式。",
        "规则引擎：覆盖命令执行、凭据泄露、路径授权、网络暴露、容器控制面和提示注入等风险。",
        "报告持久化：每次扫描生成 JSON 与 Markdown 报告，并保存到 reports 目录。",
        "误报说明：每条风险包含可能误报场景，便于人工复核和答辩说明。",
    ]
    for item in bullets:
        add_bullet(doc, item)

    add_heading(doc, "检测规则概览", 1)
    add_table(
        doc,
        [
            ("命令执行", "高风险命令、Shell 内联命令、运行时包管理器启动。"),
            ("权限暴露", "敏感目录、根目录、用户主目录、通配符路径、写入权限。"),
            ("凭据风险", "环境变量或启动参数中的 Token、Key、Secret、Password。"),
            ("网络风险", "0.0.0.0 绑定、无鉴权 HTTP、远程 HTTP(S) MCP 入口。"),
            ("容器风险", "Docker Socket 暴露。"),
            ("提示注入", "工具描述中诱导模型忽略原有指令或泄露系统提示词的文本。"),
        ],
    )

    add_heading(doc, "影响力", 1)
    add_body(
        doc,
        "该项目围绕 AI Agent 工具调用安全构建了可复现的本地审计流程，可用于实验室 AI 安全教学、"
        "MCP 工具上线前检查、个人开发环境风险排查以及 Agent 安全治理规则研究。项目不执行目标命令、"
        "不主动连接外部目标，合规边界清晰，适合作为防御型安全工具开源和二次开发。"
    )
    add_body(
        doc,
        "当前版本已实现 Web 仪表盘、CLI、自动发现、报告导出、多格式解析和 16 个单元测试，"
        "能够支撑课程演示、项目截图、README 展示和实验室项目汇总材料。"
    )

    add_heading(doc, "项目目录截图（文字版）", 1)
    directory_items = [
        "mcp_guard/：扫描器、规则引擎、评分、报告、CLI、Web 路由。",
        "templates/：首页、扫描结果页、示例页、自动发现页。",
        "static/：仪表盘样式。",
        "examples/：低/中/高风险 JSON 示例，以及 YAML/TOML 高风险示例。",
        "tests/：解析、规则、自动发现单元测试。",
        "reports/：扫描生成的 JSON 和 Markdown 报告。",
    ]
    for item in directory_items:
        add_bullet(doc, item)

    add_heading(doc, "项目说明截图", 1)
    add_body(
        doc,
        "Web 首页提供上传配置、粘贴配置、选择示例和自动扫描本机配置入口；扫描结果页展示风险等级、风险分数、"
        "Server 数量、高危问题数量、报告文件路径和风险详情；示例页可用于课堂或答辩快速演示低风险与高风险配置差异。"
    )
    add_body(
        doc,
        "CLI 示例：python -m mcp_guard scan examples/high-risk.json；"
        "自动发现示例：python -m mcp_guard discover --scan。"
    )

    add_heading(doc, "项目声明", 1)
    add_body(
        doc,
        "本项目由暨南大学网络空间安全学院学生完成，主要用于网络安全学习、AI Agent / MCP 安全审计研究和本地配置风险排查。"
        "项目仅对配置文件和工具描述进行静态分析，不执行 MCP Server 命令，不进行真实渗透测试，不主动扫描外部网络目标。"
        "项目开源后可供研究人员和学习者进行复现、规则扩展和二次开发，使用者应遵守合法授权和数据安全要求。"
    )

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("MCP-Guard 项目情况总结")
    set_font(run, size=9, color="666666")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
    print(OUT)
